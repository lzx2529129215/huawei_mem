#!/usr/bin/env python3
"""Build a lightweight, valid DOCX fixture for Perf_WPS_0040."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "samples" / "wps" / "word_0040_fixture.docx"


def set_font(run, western: str = "Calibri", east_asia: str = "Noto Sans CJK SC") -> None:
    run.font.name = western
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, "2E74B5"),
        ("Heading 2", 13, 14, 7, "2E74B5"),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    first = doc.add_paragraph()
    first.paragraph_format.space_after = Pt(10)
    run = first.add_run("Perf_WPS_0040 自动化测试首段：本段用于复制、粘贴和查找验证。")
    set_font(run)
    run.bold = True

    for chapter in range(1, 9):
        if chapter > 1:
            doc.add_page_break()
        title = "第八章 最后章节" if chapter == 8 else f"第{chapter}章 自动化测试章节"
        doc.add_heading(title, level=1)
        for paragraph_index in range(1, 7):
            paragraph = doc.add_paragraph()
            text = (
                f"这是第{chapter}章的第{paragraph_index}段测试内容。"
                "该文档用于验证WPS文字中的翻页、查找、导航、复制粘贴、插入图片和保存操作。"
                "每一页都保留清晰的章节标题，便于通过左侧导航跳转。"
            )
            body_run = paragraph.add_run(text)
            set_font(body_run)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
